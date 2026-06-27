from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches


MONTHS = {
    "01": "января", "02": "февраля", "03": "марта", "04": "апреля", "05": "мая", "06": "июня",
    "07": "июля", "08": "августа", "09": "сентября", "10": "октября", "11": "ноября", "12": "декабря",
}


TECHNICAL_FIELDS = {
    "id", "лист", "№ п/п", "номер", "фамилия", "имя", "отчество", "год выдачи диплома",
    "дата проведения аттестационной комиссии", "название итоговой работы", "название_курса",
    "название курса", "специальность", "направление", "период", "сроки обучения", "срок обучения",
    "часы", "количество часов", "объем часов", "объём часов", "дата_начала", "дата начала",
    "дата_окончания", "дата окончания", "дата_решения", "дата решения", "рег_номер",
    "рег номер", "регистрационный номер", "код", "номер диплома", "номер вкладыша",
    "квалификация", "присвоенная квалификация", "вид деятельности", "новый вид деятельности",
    "документ об образовании", "предыдущий документ", "университет", "образовательная организация",
}


def clean(value) -> str:
    return str(value or "").replace("\r", " ").strip()


def normalize(value: str) -> str:
    return re.sub(r"\s+", " ", clean(value)).strip().lower()


def get_value(student: dict, *names: str) -> str:
    normalized = {normalize(k).replace(" ", "_"): v for k, v in student.items()}
    for name in names:
        key = normalize(name).replace(" ", "_")
        if key in normalized:
            return clean(normalized[key])
    return ""


def format_russian_date(value: str) -> str:
    text = clean(value)
    if not text:
        return ""

    if any(month in text.lower() for month in MONTHS.values()):
        if text.endswith("года"):
            return text
        if re.search(r"\b\d{4}\b", text):
            return f"{text} года"
        return text

    match = re.search(r"(\d{1,2})[./-](\d{1,2})[./-](\d{4})", text)
    if match:
        day = int(match.group(1))
        month = f"{int(match.group(2)):02d}"
        year = match.group(3)
        return f"{day:02d} {MONTHS.get(month, month)} {year} года"

    return text


def date_parts(value: str) -> tuple[str, str, str]:
    text = format_russian_date(value).replace(" года", "").strip()
    match = re.search(r"(\d{1,2})\s+([А-Яа-яёЁ]+)\s+(\d{4})", text)
    if match:
        return f"{int(match.group(1)):02d}", match.group(2), match.group(3)

    match = re.search(r"(\d{1,2})[./-](\d{1,2})[./-](\d{4})", text)
    if match:
        month = MONTHS.get(f"{int(match.group(2)):02d}", match.group(2))
        return f"{int(match.group(1)):02d}", month, match.group(3)

    return "", "", ""


def get_context(student: dict) -> dict:
    last_name = get_value(student, "Фамилия", "last_name")
    first_name = get_value(student, "Имя", "first_name")
    middle_name = get_value(student, "Отчество", "middle_name")

    course = get_value(student, "Название_курса", "Название курса", "Специальность", "Направление")
    hours = get_value(student, "Часы", "Количество часов", "Объем часов", "Объём часов")
    start_date = format_russian_date(get_value(student, "Дата_начала", "Дата начала"))
    finish_date = format_russian_date(get_value(student, "Дата_окончания", "Дата окончания", "Дата проведения аттестационной комиссии"))

    period = get_value(student, "Период", "Сроки обучения", "Срок обучения")
    if not period and start_date and finish_date:
        period = f"с {start_date} по {finish_date}"
    else:
        period = period.replace("года года", "года")

    issue_date = format_russian_date(
        get_value(student, "Дата_выдачи", "Дата выдачи", "Дата_окончания", "Дата окончания", "Дата проведения аттестационной комиссии")
    )
    decision_date = format_russian_date(get_value(student, "Дата_решения", "Дата решения", "Дата проведения аттестационной комиссии")) or issue_date

    # В итоговой ведомости для дипломов может не быть номера диплома/вкладыша.
    # Поэтому не подставляем id как фальшивый номер: если номера нет в таблице, поле остаётся пустым.
    reg_number = get_value(student, "Рег_номер", "Рег номер", "Регистрационный номер", "Код", "Номер диплома")
    insert_code = get_value(student, "Код", "Номер вкладыша", "Номер диплома", "Рег_номер", "Рег номер")

    return {
        "id": get_value(student, "id") or "",
        "reg_number": reg_number,
        "insert_code": insert_code,
        "last_name": last_name,
        "first_name": first_name,
        "middle_name": middle_name,
        "full_name": f"{last_name}\n{first_name} {middle_name}".strip(),
        "full_name_one_line": f"{last_name} {first_name} {middle_name}".strip(),
        "course": course,
        "hours": hours,
        "period": period,
        "start_date": start_date,
        "finish_date": finish_date,
        "issue_date": issue_date,
        "decision_date": decision_date,
        "qualification": get_value(student, "Квалификация", "Присвоенная квалификация") or "",
        "activity": get_value(student, "Вид деятельности", "Новый вид деятельности") or course,
        "education_document": get_value(student, "Документ об образовании", "Предыдущий документ") or "диплом о высшем образовании",
        "university": get_value(student, "Университет", "Образовательная организация") or "Новгородском государственном университете\nимени Ярослава Мудрого",
    }


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, font_name="Calibri", size=11, bold=False, italic=False):
    run.font.name = font_name
    run_element = run._element
    r_pr = run_element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_cell_text(cell, text, *, align="center", font_name="Calibri", size=11, bold=False, italic=False, margins=True):
    cell.text = clean(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if margins:
        set_cell_margins(cell, 0, 0, 0, 0)

    alignments = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    for paragraph in cell.paragraphs:
        paragraph.alignment = alignments.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            set_font(run, font_name, size, bold, italic)


def safe_file_part(value: str) -> str:
    value = clean(value).replace(" ", "_").replace("\n", "_")
    for symbol in '<>:"/\\|?*\t\r':
        value = value.replace(symbol, "")
    return value or "student"


def grade_text(value: str) -> str:
    text = normalize(value)
    if "отл" in text or text == "5":
        return "отлично"
    if "хор" in text or text == "4":
        return "хорошо"
    if "уд" in text or text == "3":
        return "удовлетворительно"
    if "зач" in text:
        return "зачтено"
    return clean(value)


def extract_modules(student: dict) -> list[dict]:
    modules = []
    for key, value in student.items():
        key_clean = clean(key)
        if normalize(key_clean) in TECHNICAL_FIELDS or not clean(value):
            continue

        key_norm = normalize(key_clean)
        if key_clean.startswith("Модуль") or "аттестация" in key_norm or "экзамен" in key_norm:
            hours_match = re.search(r"(\d+)\s*(?:час|ч)", key_clean, flags=re.IGNORECASE)
            if not hours_match:
                hours_match = re.search(r",\s*(\d+)\s*$", key_clean)

            hours = hours_match.group(1) if hours_match else ""
            name = re.sub(r",?\s*\d+\s*(?:час|ч).*", "", key_clean, flags=re.IGNORECASE).strip()
            name = re.sub(r",\s*\d+\s*$", "", name).strip()
            name = re.sub(r"^Модуль\s*\d+\.?\s*", "", name, flags=re.IGNORECASE).strip()
            modules.append({"name": name, "hours": hours, "grade": grade_text(value)})

    return modules


def fill_diploma_blank_document(template_path: Path, student: dict) -> Document:
    c = get_context(student)
    doc = Document(template_path)
    left = doc.tables[1]
    main = doc.tables[2]

    set_cell_text(left.rows[0].cells[1], c["reg_number"], size=11, bold=True)
    set_cell_text(left.rows[5].cells[1], c["issue_date"], size=11, bold=True)
    set_cell_text(main.rows[2].cells[0], c["full_name"], size=14, bold=True, italic=True)
    set_cell_text(main.rows[5].cells[0], f"прошел(а) профессиональную переподготовку по программе\n{c['course']}", size=13, bold=True, italic=True)
    set_cell_text(main.rows[7].cells[0], f"в объёме {c['hours']} часов", size=11, italic=True)
    set_cell_text(main.rows[9].cells[0], c["period"], size=11, bold=True)

    qualification_text = "Итоговая аттестационная комиссия"
    if c["decision_date"]:
        qualification_text += f"\nрешением от {c['decision_date']} подтверждает"
    qualification_text += "\nприсвоение квалификации"
    if c["qualification"]:
        qualification_text += f"\n{c['qualification']}"

    set_cell_text(main.rows[11].cells[0], qualification_text, size=11)
    set_cell_text(
        main.rows[12].cells[0],
        f"и дает право на ведение нового вида профессиональной деятельности\n{c['activity']}",
        size=11,
    )

    return doc


def fill_diploma_blank(template_path: Path, output_path: Path, student: dict):
    doc = fill_diploma_blank_document(template_path, student)
    doc.save(output_path)


def fill_insert_document(template_path: Path, student: dict) -> Document:
    c = get_context(student)
    modules = extract_modules(student)
    doc = Document(template_path)
    top = doc.tables[0]
    table = doc.tables[1]

    set_cell_text(top.rows[0].cells[11], c["insert_code"], size=10)
    set_cell_text(top.rows[3].cells[7], c["last_name"], size=11, bold=True, italic=True)
    set_cell_text(top.rows[4].cells[0], f"{c['first_name']} {c['middle_name']}", size=11, bold=True, italic=True)
    set_cell_text(top.rows[7].cells[0], c["education_document"], size=10)

    start_day, start_month, start_year = date_parts(c["start_date"])
    finish_day, finish_month, finish_year = date_parts(c["finish_date"])
    for idx in [1, 2, 3, 4, 6, 9, 10, 11, 12, 13, 14]:
        if idx < len(top.rows[8].cells):
            set_cell_text(top.rows[8].cells[idx], "", size=10)
    set_cell_text(top.rows[8].cells[1], start_day, size=10)
    set_cell_text(top.rows[8].cells[2], start_month, size=10)
    set_cell_text(top.rows[8].cells[6], start_year, size=10)
    set_cell_text(top.rows[8].cells[9], finish_day, size=10)
    set_cell_text(top.rows[8].cells[10], finish_month, size=10)
    set_cell_text(top.rows[8].cells[12], finish_year, size=10)

    set_cell_text(top.rows[11].cells[0], c["university"], size=10)
    set_cell_text(top.rows[14].cells[0], c["course"], size=10, bold=True, italic=True)

    # Нижняя таблица вкладыша теперь строится по данным из Excel.
    # Если в итоговой ведомости 8 модулей + итоговая аттестация, все строки попадут во вкладыш.
    if len(table.rows) > 1:
        module_row_xml = deepcopy(table.rows[1]._tr)
        total_row_xml = deepcopy(table.rows[-1]._tr)
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        for idx, module in enumerate(modules, start=1):
            table._tbl.append(deepcopy(module_row_xml))
            row = table.rows[-1]
            for cell in row.cells:
                set_cell_text(cell, "", size=8)
            set_cell_text(row.cells[0], f"{idx}.", size=8)
            set_cell_text(row.cells[1], module["name"], align="left", size=8)
            set_cell_text(row.cells[5], module.get("hours", ""), size=8)
            set_cell_text(row.cells[7], module["grade"], size=8)

        table._tbl.append(deepcopy(total_row_xml))
        total_row = table.rows[-1]
        for cell in total_row.cells:
            set_cell_text(cell, "", size=9)
        if c["hours"]:
            set_cell_text(total_row.cells[2], f"{c['hours']} ч.", size=9, bold=True)

    return doc

def fill_insert(template_path: Path, output_path: Path, student: dict):
    doc = fill_insert_document(template_path, student)
    doc.save(output_path)


def short_period_for_statement(student: dict) -> str:
    c = get_context(student)
    start = c.get("start_date", "").replace(" года", "")
    finish = c.get("finish_date", "").replace(" года", "")
    return f"({start} – {finish})" if start and finish else ""


def ensure_table_columns(table, required_count: int):
    while len(table.rows[0].cells) < required_count:
        table.add_column(Inches(0.85))


def fill_statement(template_path: Path, output_path: Path, students: list[dict]):
    doc = Document(template_path)
    table = doc.tables[0]

    first_modules = extract_modules(students[0]) if students else []
    required_columns = 2 + len(first_modules)
    ensure_table_columns(table, required_columns)

    if len(doc.paragraphs) > 1:
        period = short_period_for_statement(students[0]) if students else ""
        if period:
            doc.paragraphs[1].text = period

    # Удаляем старые строки кроме шапки.
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    header = table.rows[0]
    set_cell_text(header.cells[0], "№", size=8, bold=True)
    set_cell_text(header.cells[1], "ФИО", size=8, bold=True)
    for col_index, module in enumerate(first_modules, start=2):
        set_cell_text(header.cells[col_index], module["name"], size=7, bold=True)

    for index, student in enumerate(students, start=1):
        c = get_context(student)
        modules = extract_modules(student)
        row = table.add_row()
        set_cell_text(row.cells[0], str(index), size=8)
        set_cell_text(row.cells[1], c["full_name_one_line"], align="left", size=8)
        for col_index in range(2, required_columns):
            module_index = col_index - 2
            value = modules[module_index]["grade"] if module_index < len(modules) else ""
            set_cell_text(row.cells[col_index], value, size=8)

    doc.save(output_path)

def _insert_before_section_properties(target_doc: Document, element):
    body = target_doc.element.body
    sect_pr = body.sectPr
    if sect_pr is not None:
        body.insert(body.index(sect_pr), element)
    else:
        body.append(element)


def append_document_pages(target_doc: Document, source_doc: Document, add_page_break=True):
    if add_page_break:
        paragraph = target_doc.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    for child in source_doc.element.body:
        if child.tag == qn("w:sectPr"):
            continue
        _insert_before_section_properties(target_doc, deepcopy(child))


def save_combined_documents(
    students: list[dict],
    diploma_template: Path,
    insert_template: Path,
    generated_folder: Path,
) -> tuple[Path, Path]:
    diploma_docs = [fill_diploma_blank_document(diploma_template, student) for student in students]
    insert_docs = [fill_insert_document(insert_template, student) for student in students]

    diploma_output = generated_folder / "Дипломы.docx"
    insert_output = generated_folder / "Вкладыши.docx"

    master_diploma = diploma_docs[0]
    for doc in diploma_docs[1:]:
        append_document_pages(master_diploma, doc, add_page_break=True)
    master_diploma.save(diploma_output)

    master_insert = insert_docs[0]
    for doc in insert_docs[1:]:
        append_document_pages(master_insert, doc, add_page_break=True)
    master_insert.save(insert_output)

    return diploma_output, insert_output


def generate_diploma_documents_zip(
    students: list[dict],
    diploma_template: Path,
    insert_template: Path,
    statement_template: Path,
    generated_folder: Path,
) -> Path:
    """Создаёт нормальный комплект для дипломов.

    Внутри архива теперь не десятки отдельных файлов, а три готовых документа:
    1) Дипломы.docx — все дипломы, каждый слушатель с новой страницы.
    2) Вкладыши.docx — все вкладыши, каждый слушатель с новой страницы.
    3) Сводная_ведомость.docx — общая ведомость по всем слушателям.
    """
    if not students:
        raise ValueError("Нет данных для формирования дипломов")

    generated_folder.mkdir(exist_ok=True)
    output_folder = generated_folder / "diploma_package"
    output_folder.mkdir(exist_ok=True)

    diploma_output, insert_output = save_combined_documents(
        students=students,
        diploma_template=diploma_template,
        insert_template=insert_template,
        generated_folder=output_folder,
    )

    statement_path = output_folder / "Сводная_ведомость.docx"
    fill_statement(statement_template, statement_path, students)

    zip_path = generated_folder / "ready_diplomas.zip"
    with ZipFile(zip_path, "w", compression=ZIP_DEFLATED) as zip_file:
        for file_path in [diploma_output, insert_output, statement_path]:
            zip_file.write(file_path, arcname=file_path.name)

    return zip_path
