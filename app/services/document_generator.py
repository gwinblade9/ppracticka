from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def get_value(student: dict, *names: str) -> str:
    """Получает значение по разным вариантам названия колонки."""
    normalized = {str(k).strip().lower().replace(" ", "_"): v for k, v in student.items()}

    for name in names:
        key = name.strip().lower().replace(" ", "_")
        if key in normalized:
            return str(normalized[key]).strip()

    return ""


def build_student_context(student: dict) -> dict:
    last_name = get_value(student, "Фамилия", "last_name")
    first_name = get_value(student, "Имя", "first_name")
    middle_name = get_value(student, "Отчество", "middle_name")

    return {
        "reg_number": get_value(student, "Рег_номер", "Рег номер", "Регистрационный номер", "Код"),
        "issue_date": get_value(student, "Дата_выдачи", "Дата выдачи"),
        "last_name": last_name,
        "first_name": first_name,
        "middle_name": middle_name,
        # ФИО оставляем в две строки, как в исходном бланке.
        "full_name": f"{last_name}\n{first_name} {middle_name}".strip(),
        "city": get_value(student, "Город", "Место_выдачи", "Место выдачи", "city"),
        "period": get_value(student, "Период", "Сроки обучения", "Срок обучения"),
        "course": get_value(student, "Название_курса", "Название курса", "Курс", "Специальность", "Направление"),
        "hours": get_value(student, "Часы", "Количество часов", "Объем часов", "Объём часов"),
    }


def clean_text(value: str) -> str:
    """Убирает пробелы по краям, но сохраняет переносы строк внутри ФИО."""
    return str(value or "").strip()


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0) -> None:
    """Убирает внутренние отступы ячейки, чтобы текст реально был по центру."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, *, font_name="Calibri", font_size=None, bold=None, italic=None) -> None:
    """Настраивает шрифт. qn нужен, чтобы кириллица тоже была Calibri."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)

    if font_size is not None:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_text(
    cell,
    text: str,
    *,
    align: str = "left",
    font_name: str = "Calibri",
    font_size: int | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    remove_cell_margins: bool = True,
) -> None:
    """Записывает текст в ячейку и настраивает выравнивание/шрифт."""
    cell.text = clean_text(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if remove_cell_margins:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)

        for run in paragraph.runs:
            set_run_font(
                run,
                font_name=font_name,
                font_size=font_size,
                bold=bold,
                italic=italic,
            )


def fill_template(template_path: Path, output_path: Path, student: dict) -> None:
    """Подставляет данные студента в конкретные ячейки бланка."""
    context = build_student_context(student)
    document = Document(template_path)

    # В файле есть 3 таблицы. Нумерация в Python начинается с 0.
    left_block = document.tables[1]
    main_block = document.tables[2]

    # Уникальный номер / код — строго по центру.
    set_cell_text(
        left_block.rows[1].cells[0],
        context["reg_number"],
        align="center",
        font_name="Calibri",
        font_size=11,
        bold=True,
        italic=False,
        remove_cell_margins=True,
    )

    # Дата выдачи — по центру.
    set_cell_text(
        left_block.rows[6].cells[0],
        context["issue_date"],
        align="center",
        font_name="Calibri",
        font_size=11,
        bold=True,
        italic=False,
    )

    # Город — без ведущих пробелов и строго по центру.
    set_cell_text(
        left_block.rows[4].cells[0],
        context["city"] or left_block.rows[4].cells[0].text,
        align="center",
        font_name="Calibri",
        font_size=11,
        bold=True,
        italic=False,
    )

    # ФИО — Calibri 14, курсив, по центру.
    set_cell_text(
        main_block.rows[3].cells[0],
        context["full_name"],
        align="center",
        font_name="Calibri",
        font_size=14,
        bold=True,
        italic=True,
    )

    # Сроки обучения — Calibri 11, жирный, по центру.
    set_cell_text(
        main_block.rows[8].cells[0],
        context["period"],
        align="center",
        font_name="Calibri",
        font_size=11,
        bold=True,
        italic=False,
    )

    # Наименование специальности / курса — Calibri 13, жирный курсив, по центру.
    set_cell_text(
        main_block.rows[12].cells[0],
        context["course"],
        align="center",
        font_name="Calibri",
        font_size=13,
        bold=True,
        italic=True,
    )

    # Объём часов — Calibri 11, курсив, по центру.
    set_cell_text(
        main_block.rows[16].cells[0],
        f"в объеме {context['hours']} ч.",
        align="center",
        font_name="Calibri",
        font_size=11,
        bold=False,
        italic=True,
    )

    document.save(output_path)


def safe_file_part(value: str) -> str:
    value = value.strip().replace(" ", "_")
    bad_symbols = '<>:"/\\|?*\n\r\t'
    for symbol in bad_symbols:
        value = value.replace(symbol, "")
    return value or "student"


def insert_before_final_section(document: Document, element) -> None:
    body = document.element.body
    section = body.sectPr

    if section is None:
        body.append(element)
        return

    body.insert(body.index(section), element)


def append_page_break(document: Document) -> None:
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")

    run.append(page_break)
    paragraph.append(run)
    insert_before_final_section(document, paragraph)


def append_document_body(target: Document, source: Document) -> None:
    for element in source.element.body:
        if element.tag == qn("w:sectPr"):
            continue
        insert_before_final_section(target, deepcopy(element))


def generate_documents_docx(students: list[dict], template_path: Path, generated_folder: Path) -> Path:
    generated_folder.mkdir(exist_ok=True)

    output_doc = generated_folder / "certificates.docx"
    temp_file = generated_folder / "_temp_certificate.docx"
    merged_document = None

    for student in students:
        fill_template(template_path, temp_file, student)
        current = Document(temp_file)

        if merged_document is None:
            merged_document = current
        else:
            append_page_break(merged_document)
            append_document_body(merged_document, current)

    if merged_document is None:
        merged_document = Document()

    merged_document.save(output_doc)

    try:
        temp_file.unlink()
    except FileNotFoundError:
        pass

    return output_doc



def generate_documents_zip(students: list[dict], template_path: Path, generated_folder: Path) -> Path:
    return generate_documents_docx(students, template_path, generated_folder)
