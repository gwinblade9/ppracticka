from pathlib import Path
import re
from docx import Document


MONTHS = {
    "01": "января",
    "02": "февраля",
    "03": "марта",
    "04": "апреля",
    "05": "мая",
    "06": "июня",
    "07": "июля",
    "08": "августа",
    "09": "сентября",
    "10": "октября",
    "11": "ноября",
    "12": "декабря",
}


def clean_cell_text(value: str) -> str:
    """Приводит текст из ячейки Word к нормальному виду."""
    return re.sub(r"\s+", " ", str(value or "").replace("\xa0", " ")).strip()


def date_to_words(value: str) -> str:
    """Преобразует 05.06.2026 в 05 июня 2026 года. Если формат другой — возвращает как есть."""
    value = clean_cell_text(value)
    match = re.fullmatch(r"(\d{1,2})[.](\d{1,2})[.](\d{4})", value)
    if not match:
        return value

    day, month, year = match.groups()
    return f"{int(day):02d} {MONTHS.get(month.zfill(2), month)} {year} года"


def extract_statement_meta(document: Document) -> dict:
    """Достаёт программу, период и часы из шапки ведомости."""
    paragraphs = [clean_cell_text(paragraph.text) for paragraph in document.paragraphs]
    paragraphs = [text for text in paragraphs if text]
    full_text = "\n".join(paragraphs)

    course = ""
    period = ""
    hours = ""
    end_date = ""

    # Программа обычно идёт после строки "Наименование программы повышения квалификации:".
    for index, text in enumerate(paragraphs):
        if "Наименование программы" in text:
            course_parts = []
            for next_text in paragraphs[index + 1:]:
                if next_text.lower().startswith("период обучения"):
                    break
                course_parts.append(next_text)
            course = " ".join(course_parts).strip(" «»\"")
            break

    period_match = re.search(
        r"период\s+обучения\s*:\s*с\s*(\d{1,2}\.\d{1,2}\.\d{4})\s*по\s*(\d{1,2}\.\d{1,2}\.\d{4})",
        full_text,
        flags=re.IGNORECASE,
    )
    if period_match:
        start_raw, end_raw = period_match.groups()
        start_words = date_to_words(start_raw)
        end_words = date_to_words(end_raw)
        period = f"с {start_words} по {end_words}"
        end_date = end_words

    hours_match = re.search(r"об[ъь]ем\s*(\d+)\s*час", full_text, flags=re.IGNORECASE)
    if hours_match:
        hours = hours_match.group(1)

    return {
        "course": course,
        "period": period,
        "hours": hours,
        "issue_date": end_date,
    }


def split_fio(full_name: str) -> tuple[str, str, str]:
    parts = clean_cell_text(full_name).split()
    last_name = parts[0] if len(parts) >= 1 else ""
    first_name = parts[1] if len(parts) >= 2 else ""
    middle_name = " ".join(parts[2:]) if len(parts) >= 3 else ""
    return last_name, first_name, middle_name


def is_certificate_statement(document: Document) -> bool:
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return "ВЕДОМОСТЬ" in full_text and "выдачи удостоверений" in full_text


def read_certificate_statement_docx(file_path: Path) -> list[dict]:
    """Читает ведомость выдачи удостоверений.

    Важно: колонка "Номер бланка документа" специально игнорируется и никуда не сохраняется.
    Берём только: № регистрационный, ФИО, дата выдачи, программу, период и часы.
    """
    document = Document(file_path)

    if not document.tables:
        raise ValueError("В ведомости не найдена таблица со слушателями")

    meta = extract_statement_meta(document)
    table = document.tables[0]
    students = []

    # В ведомости столбцы обычно такие:
    # 0 — № п/п, 1 — № регистрац., 2 — Номер бланка документа, 3 — ФИО, 4 — Дата выдачи.
    for row in table.rows[1:]:
        cells = [clean_cell_text(cell.text) for cell in row.cells]
        if len(cells) < 4:
            continue

        reg_number = cells[1] if len(cells) > 1 else ""
        full_name = cells[3] if len(cells) > 3 else ""
        issue_date = cells[4] if len(cells) > 4 else ""

        if not reg_number and not full_name:
            continue

        last_name, first_name, middle_name = split_fio(full_name)

        students.append({
            "id": len(students) + 1,
            "Рег_номер": reg_number,
            # "Номер бланка документа" НЕ добавляем специально.
            "Дата_выдачи": date_to_words(issue_date) if issue_date else meta["issue_date"],
            "Фамилия": last_name,
            "Имя": first_name,
            "Отчество": middle_name,
            "Период": meta["period"],
            "Название_курса": meta["course"],
            "Часы": meta["hours"],
        })

    if not students:
        raise ValueError("В ведомости не найдены строки со слушателями")

    return students


def read_docx_table(file_path: Path) -> list[dict]:
    """Читает данные из Word-файла.

    Если это ведомость выдачи удостоверений, используется специальный разбор.
    Иначе читается первая таблица: первая строка считается заголовками.
    """
    document = Document(file_path)

    if is_certificate_statement(document):
        return read_certificate_statement_docx(file_path)

    if not document.tables:
        raise ValueError("В Word-файле не найдена таблица с данными")

    table = document.tables[0]
    headers = [clean_cell_text(cell.text) for cell in table.rows[0].cells]

    students = []
    for row in table.rows[1:]:
        values = [clean_cell_text(cell.text) for cell in row.cells]

        if not any(values):
            continue

        student = {headers[i]: values[i] if i < len(values) else "" for i in range(len(headers))}
        student["id"] = len(students) + 1
        students.append(student)

    return students
